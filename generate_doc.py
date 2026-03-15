"""
WAK Solutions - Solution Design Document Generator
Produces WAK_Solution_Design.docx in the project root.
"""

import os
import io
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ─────────────────────────────────────────────────────────────
C_PRIMARY   = "#0F510F"   # dark green  – headings / titles
C_ACCENT    = "#408440"   # light green – diagram accents
C_BODY      = "#4D4D4D"   # dark grey   – body text / tables
C_WHITE     = "#FFFFFF"

def hex_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

PRIMARY_RGB = hex_rgb(C_PRIMARY)
ACCENT_RGB  = hex_rgb(C_ACCENT)
BODY_RGB    = hex_rgb(C_BODY)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def save_fig(fig, dpi=150):
    """Save a matplotlib figure to a BytesIO buffer and return it."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def draw_box(ax, x, y, w, h, label, sublabel=None,
             fc=C_PRIMARY, ec=C_ACCENT, tc="white",
             fontsize=9, radius=0.04):
    """Draw a rounded rectangle with centred label."""
    box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle=f"round,pad=0.01,rounding_size={radius}",
                         linewidth=1.5, edgecolor=ec, facecolor=fc, zorder=3)
    ax.add_patch(box)
    if sublabel:
        ax.text(x, y + h*0.12, label, ha="center", va="center",
                color=tc, fontsize=fontsize, fontweight="bold", zorder=4)
        ax.text(x, y - h*0.22, sublabel, ha="center", va="center",
                color=tc, fontsize=fontsize - 1.5, style="italic", zorder=4)
    else:
        ax.text(x, y, label, ha="center", va="center",
                color=tc, fontsize=fontsize, fontweight="bold", zorder=4)


def arrow(ax, x0, y0, x1, y1, label="", color=C_ACCENT, lw=1.6):
    ax.annotate("", xy=(x1, y1), xytext=(x0, y0),
                arrowprops=dict(arrowstyle="-|>", color=color,
                                lw=lw, mutation_scale=14), zorder=2)
    if label:
        mx, my = (x0+x1)/2, (y0+y1)/2
        ax.text(mx+0.01, my+0.01, label, ha="center", va="bottom",
                fontsize=7.5, color=C_BODY, zorder=5)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 1 – System Architecture
# ══════════════════════════════════════════════════════════════════════════════

def make_architecture_diagram():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # ── swim-lane backgrounds ────────────────────────────────────────────────
    def lane(x0, y0, w, h, label, lc):
        rect = mpatches.FancyBboxPatch((x0, y0), w, h,
                                       boxstyle="round,pad=0.05",
                                       linewidth=1.2, edgecolor=lc,
                                       facecolor=lc + "18", zorder=0)
        ax.add_patch(rect)
        ax.text(x0 + w/2, y0 + h - 0.18, label, ha="center", va="top",
                fontsize=8, color=lc, fontweight="bold")

    lane(0.15, 0.2,  2.3, 5.9, "Customer",      C_BODY)
    lane(2.65, 0.2,  3.0, 5.9, "AI Service\n(Python / Railway)", C_ACCENT)
    lane(5.85, 0.2,  3.0, 5.9, "Dashboard\n(Node.js / Railway)", C_PRIMARY)
    lane(9.05, 0.2,  1.8, 5.9, "External\nAPIs",  C_BODY)

    # ── nodes ────────────────────────────────────────────────────────────────
    # Customer
    draw_box(ax, 1.3, 5.2, 1.8, 0.65, "Customer", "WhatsApp",
             fc=C_BODY, ec=C_ACCENT, tc="white")
    draw_box(ax, 1.3, 1.0, 1.8, 0.65, "Human Agent", "Browser / Mobile",
             fc=C_BODY, ec=C_ACCENT, tc="white")

    # AI Service
    draw_box(ax, 4.15, 5.2, 2.4, 0.65, "Webhook Receiver",
             "FastAPI /webhook", fc=C_ACCENT, ec=C_PRIMARY, tc="white")
    draw_box(ax, 4.15, 3.5, 2.4, 0.65, "AI Agent",
             "GPT-4.1-mini + Tools", fc=C_PRIMARY, ec=C_ACCENT, tc="white")
    draw_box(ax, 4.15, 1.9, 2.4, 0.65, "Message Sender",
             "FastAPI /send", fc=C_ACCENT, ec=C_PRIMARY, tc="white")

    # Dashboard
    draw_box(ax, 7.35, 5.2, 2.4, 0.65, "Notification Service",
             "Push + Escalations", fc=C_PRIMARY, ec=C_ACCENT, tc="white")
    draw_box(ax, 7.35, 3.5, 2.4, 0.65, "Agent Dashboard",
             "React SPA", fc=C_ACCENT, ec=C_PRIMARY, tc="white")
    draw_box(ax, 7.35, 1.9, 2.4, 0.65, "REST API",
             "Express / Node.js", fc=C_PRIMARY, ec=C_ACCENT, tc="white")

    # Shared DB — spans centre
    draw_box(ax, 5.5, 0.72, 2.6, 0.65, "PostgreSQL Database",
             "Neon (Managed Cloud)", fc=C_PRIMARY, ec=C_ACCENT, tc="white")

    # External APIs
    draw_box(ax, 9.95, 5.2, 1.4, 0.65, "Meta\nWhatsApp\nCloud API",
             fc=C_BODY, ec=C_ACCENT, tc="white", fontsize=7.5)
    draw_box(ax, 9.95, 3.5, 1.4, 0.65, "OpenAI\nGPT-4.1-mini",
             fc=C_BODY, ec=C_ACCENT, tc="white", fontsize=7.5)

    # ── arrows ───────────────────────────────────────────────────────────────
    # Customer → Webhook (inbound)
    arrow(ax, 2.2, 5.2, 2.95, 5.2, "WhatsApp msg")
    # Webhook → AI Agent
    arrow(ax, 4.15, 4.87, 4.15, 4.17, "parse & route")
    # AI Agent ↔ OpenAI
    arrow(ax, 5.35, 3.5, 9.25, 3.5, "API request")
    arrow(ax, 9.25, 3.3, 5.35, 3.3, "response")
    # AI Agent → Message Sender
    arrow(ax, 4.15, 3.17, 4.15, 2.23, "reply text")
    # Message Sender → WhatsApp API
    arrow(ax, 5.35, 1.9, 9.25, 1.9)
    # WhatsApp API → Customer
    arrow(ax, 9.25, 5.35, 9.95, 5.35)
    # Customer ← WhatsApp API
    arrow(ax, 9.25, 5.05, 5.35, 5.05)

    # Notification
    arrow(ax, 5.35, 5.2, 6.15, 5.2, "event notify")
    # Dashboard → REST API
    arrow(ax, 7.35, 3.17, 7.35, 2.23, "API calls")
    # Agent → Dashboard
    arrow(ax, 2.2, 1.0, 6.15, 1.0, "manage conversations")
    # REST API ↔ DB
    arrow(ax, 6.65, 1.57, 6.15, 1.05)
    # AI ↔ DB
    arrow(ax, 4.15, 1.57, 4.85, 1.05)

    ax.set_title("End-to-End System Architecture", fontsize=13,
                 fontweight="bold", color=C_PRIMARY, pad=12)
    return save_fig(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 2 – The Three Core Flows (side-by-side)
# ══════════════════════════════════════════════════════════════════════════════

def flow_col(ax, cx, steps, title, title_color=C_PRIMARY):
    """Draw a vertical flow column."""
    TOP = 6.8
    GAP = 1.05
    ax.text(cx, TOP + 0.25, title, ha="center", va="bottom",
            fontsize=9.5, fontweight="bold", color=title_color)
    for i, (lbl, sub, fc) in enumerate(steps):
        y = TOP - i * GAP
        draw_box(ax, cx, y, 1.9, 0.62, lbl, sub,
                 fc=fc, ec=C_ACCENT, tc="white", fontsize=8.5)
        if i < len(steps) - 1:
            arrow(ax, cx, y - 0.31, cx, y - 0.73)


def make_flows_diagram():
    fig, ax = plt.subplots(figsize=(11, 7.5))
    ax.set_xlim(0, 11); ax.set_ylim(0.3, 7.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Dividers
    for xd in [3.7, 7.4]:
        ax.axvline(xd, color="#DDDDDD", lw=1.2, zorder=0)

    # Flow 1: Customer → AI Reply
    flow_col(ax, 1.85, [
        ("Customer Sends\nMessage",     "WhatsApp",          C_BODY),
        ("Webhook Received",            "FastAPI service",   C_ACCENT),
        ("AI Processes\nMessage",       "GPT-4.1-mini",      C_PRIMARY),
        ("Order Lookup\n(if needed)",   "Database query",    C_ACCENT),
        ("Reply Generated",             "Natural language",  C_PRIMARY),
        ("Reply Delivered",             "WhatsApp → Customer", C_BODY),
    ], "① Customer → AI Reply")

    # Flow 2: Escalation
    flow_col(ax, 5.5, [
        ("AI Detects\nComplex Issue",   "Cannot resolve",    C_PRIMARY),
        ("Escalation Created",          "Saved to database", C_ACCENT),
        ("Dashboard Flagged",           "Alert to agent",    C_PRIMARY),
        ("Agent Reviews\nConversation", "Full history shown",C_ACCENT),
        ("Agent Closes\nEscalation",    "Marks resolved",    C_PRIMARY),
    ], "② AI Escalates to Agent")

    # Flow 3: Agent Reply
    flow_col(ax, 9.15, [
        ("Agent Types\nReply",          "Dashboard UI",      C_ACCENT),
        ("API Request\nSent",           "Express REST API",  C_PRIMARY),
        ("Forwarded to\nAI Service",    "Internal webhook",  C_ACCENT),
        ("Delivered via\nWhatsApp",     "Meta Cloud API",    C_BODY),
        ("Saved to\nDatabase",          "Full audit trail",  C_PRIMARY),
    ], "③ Agent Replies Manually")

    ax.set_title("The Three Core Conversation Flows", fontsize=13,
                 fontweight="bold", color=C_PRIMARY, pad=10)
    return save_fig(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 3 – Dashboard Wireframe
# ══════════════════════════════════════════════════════════════════════════════

def make_dashboard_wireframe():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Chrome border
    chrome = FancyBboxPatch((0.2, 0.2), 10.6, 6.1,
                            boxstyle="round,pad=0.05",
                            linewidth=2, edgecolor=C_PRIMARY,
                            facecolor="white", zorder=1)
    ax.add_patch(chrome)

    # Header bar
    hdr = FancyBboxPatch((0.2, 5.8), 10.6, 0.5,
                         boxstyle="round,pad=0",
                         linewidth=0, facecolor=C_PRIMARY, zorder=2)
    ax.add_patch(hdr)
    ax.text(1.0, 6.05, "WAK Solutions — Agent Dashboard",
            ha="left", va="center", fontsize=10, fontweight="bold",
            color="white", zorder=3)
    ax.text(10.5, 6.05, "[x]  Logout",
            ha="right", va="center", fontsize=8.5, color="white", zorder=3)

    # Sidebar
    side = FancyBboxPatch((0.2, 0.2), 2.6, 5.6,
                          boxstyle="round,pad=0",
                          linewidth=0, facecolor="#F5F8F5", zorder=2)
    ax.add_patch(side)
    ax.axvline(2.8, color=C_ACCENT, lw=1, zorder=3)
    ax.text(1.5, 5.65, "Conversations", ha="center", va="center",
            fontsize=8.5, fontweight="bold", color=C_PRIMARY, zorder=4)

    # Conversation rows in sidebar
    conv_data = [
        ("+971 50 123 4567", "My order hasn't arrived...", True, False),
        ("+86 138 0000 1234", "Where is my order?",         False, False),
        ("+971 52 987 6543", "I need to speak to someone", False, True),
        ("+971 55 444 3333", "What time do you close?",   False, False),
        ("+44 7700 900 123", "Can I get a refund?",        False, False),
    ]
    for i, (phone, msg, selected, escalated) in enumerate(conv_data):
        y_row = 5.2 - i * 0.88
        bg_col = "#E8F2E8" if selected else "white"
        row_bg = FancyBboxPatch((0.25, y_row - 0.35), 2.5, 0.72,
                                boxstyle="round,pad=0.02",
                                linewidth=0.5,
                                edgecolor=C_ACCENT if selected else "#EEEEEE",
                                facecolor=bg_col, zorder=3)
        ax.add_patch(row_bg)
        ax.text(0.45, y_row + 0.16, phone, ha="left", va="center",
                fontsize=7.2, fontweight="bold", color=C_PRIMARY, zorder=4)
        ax.text(0.45, y_row - 0.10, msg[:32] + ("…" if len(msg) > 32 else ""),
                ha="left", va="center", fontsize=6.5, color=C_BODY, zorder=4)
        if escalated:
            badge = FancyBboxPatch((2.3, y_row + 0.05), 0.38, 0.22,
                                   boxstyle="round,pad=0.02",
                                   linewidth=0, facecolor="#C0392B", zorder=5)
            ax.add_patch(badge)
            ax.text(2.49, y_row + 0.16, "OPEN", ha="center", va="center",
                    fontsize=5.5, color="white", fontweight="bold", zorder=6)

    # Chat area
    ax.text(5.5, 5.65, "+971 50 123 4567", ha="center", va="center",
            fontsize=9, fontweight="bold", color=C_PRIMARY, zorder=4)

    # Chat bubbles
    bubbles = [
        (3.0, 4.9, 2.6, 0.55, "My order hasn't arrived yet.\nOrder #ORD-2024-001",
         "white", C_BODY, True, "Customer · 10:31 AM"),
        (5.8, 4.1, 3.1, 0.65,
         "Hello! I found your order #ORD-2024-001.\nIt's currently out for delivery today.",
         "#E8F2E8", C_PRIMARY, False, "AI Agent · 10:31 AM"),
        (3.0, 3.2, 2.5, 0.55, "How long will it take?",
         "white", C_BODY, True, "Customer · 10:32 AM"),
        (5.8, 2.5, 3.0, 0.55, "You can expect delivery between\n2 PM and 6 PM today.",
         "#E8F2E8", C_PRIMARY, False, "AI Agent · 10:32 AM"),
    ]
    for bx, by, bw, bh, btxt, bfc, btc, is_left, ts in bubbles:
        b = FancyBboxPatch((bx, by - bh/2), bw, bh,
                           boxstyle="round,pad=0.06",
                           linewidth=0.8, edgecolor=C_ACCENT,
                           facecolor=bfc, zorder=3)
        ax.add_patch(b)
        ax.text(bx + 0.12, by, btxt, ha="left", va="center",
                fontsize=6.8, color=btc, zorder=4)
        ax.text(bx if is_left else bx + bw, by - bh/2 - 0.12,
                ts, ha="left" if is_left else "right",
                va="top", fontsize=5.8, color="#999999", zorder=4)

    # Reply input box
    input_box = FancyBboxPatch((2.9, 0.32), 6.2, 0.52,
                               boxstyle="round,pad=0.04",
                               linewidth=1, edgecolor=C_ACCENT,
                               facecolor="white", zorder=3)
    ax.add_patch(input_box)
    ax.text(3.1, 0.58, "Type a reply…",
            ha="left", va="center", fontsize=7.5, color="#AAAAAA", zorder=4)

    send_btn = FancyBboxPatch((9.2, 0.32), 1.4, 0.52,
                              boxstyle="round,pad=0.04",
                              linewidth=0, facecolor=C_PRIMARY, zorder=3)
    ax.add_patch(send_btn)
    ax.text(9.9, 0.58, "Send ▶",
            ha="center", va="center", fontsize=8, color="white",
            fontweight="bold", zorder=4)

    # Callout annotations
    annots = [
        (0.5, -0.1, 1.5, 0.1, "Conversation List\n(live-updating)"),
        (5.5, -0.1, 5.5, 0.1, "Message Thread\n(3-second refresh)"),
        (9.9, -0.1, 9.9, 0.1, "Send Reply\nto Customer"),
    ]
    for ax_x, ay, tx, ty, lbl in annots:
        ax.annotate(lbl, xy=(ax_x, ay + 0.2), xytext=(tx, ty - 0.05),
                    ha="center", va="top", fontsize=6.5, color=C_BODY,
                    arrowprops=dict(arrowstyle="-", color=C_ACCENT, lw=0.8))

    ax.set_title("Agent Dashboard — Annotated View", fontsize=13,
                 fontweight="bold", color=C_PRIMARY, pad=10)
    return save_fig(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 4 – Deployment Topology
# ══════════════════════════════════════════════════════════════════════════════

def make_topology_diagram():
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0.3, 4.8)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Cloud background
    cloud_bg = FancyBboxPatch((0.3, 0.5), 9.4, 3.8,
                              boxstyle="round,pad=0.1",
                              linewidth=1.5, edgecolor="#CCCCCC",
                              facecolor="#FAFAFA", zorder=0)
    ax.add_patch(cloud_bg)
    ax.text(5.0, 4.25, "Cloud Infrastructure", ha="center", va="center",
            fontsize=9, color=C_BODY, style="italic")

    nodes = [
        (1.5,  2.4, "Customer\nDevices",           "WhatsApp",       C_BODY,    1.8, 0.9),
        (4.0,  3.5, "AI Service",                  "Python FastAPI\nRailway",  C_PRIMARY, 1.9, 0.9),
        (4.0,  1.3, "Agent Dashboard",             "Node.js + React\nRailway", C_ACCENT,  1.9, 0.9),
        (7.0,  2.4, "PostgreSQL",                  "Neon Managed DB",C_PRIMARY, 1.9, 0.9),
        (9.3,  3.5, "Meta\nWhatsApp\nCloud API",   "",               C_BODY,    1.2, 0.9),
        (9.3,  1.3, "OpenAI\nGPT-4.1-mini",       "",               C_BODY,    1.2, 0.9),
    ]
    for (x, y, lbl, sub, fc, w, h) in nodes:
        draw_box(ax, x, y, w, h, lbl, sub if sub else None,
                 fc=fc, ec=C_ACCENT, tc="white", fontsize=8.5)

    connections = [
        (1.5, 2.65, 3.05, 3.5,  "Inbound msgs"),
        (1.5, 2.15, 3.05, 1.3,  "Agent access"),
        (4.95, 3.5, 6.05, 2.65, "Read/write"),
        (4.95, 1.3, 6.05, 2.15, "Read/write"),
        (4.95, 3.5, 8.7,  3.5,  "API calls"),
        (8.7,  1.3, 4.95, 1.3,  ""),
        (8.7,  3.3, 4.95, 3.3,  ""),
        (9.3,  3.05, 9.3, 1.75, ""),
        (4.0,  2.95, 4.0, 1.75, "Internal\nwebhook"),
    ]
    for (x0, y0, x1, y1, lbl) in connections:
        arrow(ax, x0, y0, x1, y1, lbl)

    ax.set_title("Deployment Topology", fontsize=13,
                 fontweight="bold", color=C_PRIMARY, pad=10)
    return save_fig(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def set_run_color(run, hex_color):
    r, g, b = hex_rgb(hex_color)
    run.font.color.rgb = RGBColor(r, g, b)


def set_cell_bg(cell, hex_color):
    """Fill a table cell background."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders(cell, hex_color="408440"):
    """Set all borders of a cell to a given colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), hex_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*PRIMARY_RGB)
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(13)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(*BODY_RGB)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    return p


def add_image(doc, buf, width_inches=6.2):
    doc.add_picture(buf, width=Inches(width_inches))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_PRIMARY.lstrip("#"))
        set_cell_borders(cell, C_ACCENT.lstrip("#"))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "#F5F8F5" if r_idx % 2 == 0 else "#FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg.lstrip("#"))
            set_cell_borders(cell, C_ACCENT.lstrip("#"))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.color.rgb = RGBColor(*BODY_RGB)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_page_break(doc):
    doc.add_page_break()


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(*BODY_RGB)
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.italic = True


def add_footer_page_numbers(doc):
    """Add page numbers to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*BODY_RGB)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*BODY_RGB)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_document(output_path):
    doc = Document()
    set_margins(doc)
    add_footer_page_numbers(doc)

    # ─── COVER PAGE ───────────────────────────────────────────────────────────
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    # Logo (if present)
    logo_path = os.path.join(os.path.dirname(__file__),
                             "wak-dash", "client", "public", "logo.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.8))

    doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("WAK Solutions")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*PRIMARY_RGB)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("WhatsApp AI Chatbot")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(22)
    sub_run.font.color.rgb = RGBColor(*ACCENT_RGB)

    doc.add_paragraph("")
    doc2_p = doc.add_paragraph()
    doc2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc2_run = doc2_p.add_run("Solution Design Document")
    doc2_run.font.name = "Calibri"
    doc2_run.font.size = Pt(14)
    doc2_run.font.color.rgb = RGBColor(*BODY_RGB)

    doc.add_paragraph("")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("March 2026  |  Version 1.0  |  Confidential")
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(*BODY_RGB)
    date_run.font.italic = True

    add_page_break(doc)

    # ─── PAGE 1: EXECUTIVE SUMMARY ────────────────────────────────────────────
    add_heading(doc, "1.  Executive Summary")
    add_body(doc,
        "WAK Solutions has built a WhatsApp-first customer service platform that "
        "uses artificial intelligence to handle customer enquiries automatically — "
        "around the clock, in multiple languages — without requiring a human agent "
        "to be present for every interaction."
    )
    doc.add_paragraph("")
    add_body(doc,
        "Customers interact through the WhatsApp application they already use every "
        "day. The AI agent understands their questions, looks up live order data "
        "from the company database, and responds instantly in the customer's preferred "
        "language. When a situation is too complex for the AI to handle, the system "
        "automatically escalates the conversation to a human agent, who can take over "
        "seamlessly from a dedicated web dashboard."
    )
    doc.add_paragraph("")
    add_heading(doc, "Key Business Outcomes", level=2)
    bullets = [
        "24/7 customer support with no additional staffing cost for routine enquiries.",
        "Instant order tracking: customers receive real-time delivery status without waiting in a queue.",
        "Multilingual coverage: the AI responds natively in Arabic, English, Chinese, and other languages.",
        "Escalation safety net: any conversation that exceeds AI capability is immediately flagged for human review.",
        "Full audit trail: every message — AI or human — is stored and searchable.",
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_paragraph("")
    add_heading(doc, "Document Scope", level=2)
    add_body(doc,
        "This document provides a non-technical overview of the system for business "
        "and management stakeholders. It covers what the system does, how its components "
        "fit together, the flows a customer and agent experience, and the infrastructure "
        "that underpins it. Implementation-level code details are intentionally excluded."
    )

    add_page_break(doc)

    # ─── PAGE 2: SYSTEM ARCHITECTURE ─────────────────────────────────────────
    add_heading(doc, "2.  How the System Works")
    add_body(doc,
        "The platform is composed of three integrated components that work together "
        "in real time: an AI Service that talks to WhatsApp, an Agent Dashboard that "
        "human staff use, and a shared database that keeps both in sync."
    )
    doc.add_paragraph("")
    arch_buf = make_architecture_diagram()
    add_image(doc, arch_buf, width_inches=6.5)
    add_caption(doc, "Figure 1 — End-to-End System Architecture")
    doc.add_paragraph("")

    add_heading(doc, "Component Summary", level=2)
    add_table(doc,
        ["Component", "Role", "Technology"],
        [
            ["AI Service",        "Receives WhatsApp messages, calls the AI, and sends replies back to the customer.", "Python / FastAPI on Railway"],
            ["Agent Dashboard",   "Web interface where human agents monitor conversations and reply manually when needed.", "React / Node.js on Railway"],
            ["PostgreSQL Database","Central store for every message, conversation, and escalation record.",              "Neon Managed Cloud"],
            ["Meta WhatsApp API", "The channel through which customers send and receive messages.",                      "Meta Cloud API v19.0"],
            ["OpenAI GPT-4.1-mini","The AI model that understands customer messages and generates natural replies.",     "OpenAI Cloud API"],
        ],
        col_widths=[1.5, 3.2, 2.0]
    )

    add_page_break(doc)

    # ─── PAGE 3: THREE CORE FLOWS ────────────────────────────────────────────
    add_heading(doc, "3.  The Three Core Flows")
    add_body(doc,
        "Every interaction in the system follows one of three paths. The diagram below "
        "shows how each flow progresses from start to finish."
    )
    doc.add_paragraph("")
    flows_buf = make_flows_diagram()
    add_image(doc, flows_buf, width_inches=6.5)
    add_caption(doc, "Figure 2 — The Three Core Conversation Flows")
    doc.add_paragraph("")

    add_heading(doc, "Flow Descriptions", level=2)
    add_table(doc,
        ["Flow", "Trigger", "Outcome"],
        [
            ["① Customer → AI Reply",
             "Customer sends any WhatsApp message to the business number.",
             "AI reads history, queries the database if needed, and sends an instant reply via WhatsApp."],
            ["② AI Escalates to Agent",
             "AI cannot resolve the issue (e.g. a complaint or sensitive request).",
             "Escalation record is created, the dashboard flags the conversation, and an agent is notified."],
            ["③ Agent Replies Manually",
             "Agent types a response in the dashboard after reviewing the conversation.",
             "Message is delivered to the customer via WhatsApp and saved to the database."],
        ],
        col_widths=[1.5, 2.7, 3.0]
    )

    add_page_break(doc)

    # ─── PAGE 4: AGENT DASHBOARD ─────────────────────────────────────────────
    add_heading(doc, "4.  The Agent Dashboard")
    add_body(doc,
        "The Agent Dashboard is a web application that human agents open in any "
        "browser or install as a mobile app. It gives agents a real-time view of "
        "all active conversations and the tools they need to respond."
    )
    doc.add_paragraph("")
    dash_buf = make_dashboard_wireframe()
    add_image(doc, dash_buf, width_inches=6.5)
    add_caption(doc, "Figure 3 — Agent Dashboard Annotated View")
    doc.add_paragraph("")

    add_heading(doc, "Key Features", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["Conversation List",      "Left sidebar shows all active conversations sorted by most recent activity. Escalated conversations are highlighted with a red badge."],
            ["Message Thread",         "Clicking any conversation opens the full chat history, showing messages from the customer, the AI, and previous agent replies."],
            ["Agent Reply Box",        "Agents type a reply at the bottom and click Send. The message is delivered to the customer's WhatsApp immediately."],
            ["Push Notifications",     "The browser or mobile app sends a notification when a new message arrives or an escalation is opened — even when the dashboard tab is not active."],
            ["Biometric Login",        "Agents can register their device fingerprint or Face ID to log in securely without typing a password each time."],
            ["Close Escalation",       "Once an agent has resolved an escalated conversation, they mark it closed. The dashboard removes the alert and the record is archived."],
            ["Mobile / PWA Support",   "The dashboard can be installed as a Progressive Web App on any smartphone, giving agents full functionality on the go."],
        ],
        col_widths=[2.0, 5.2]
    )

    add_page_break(doc)

    # ─── PAGE 5: AI CAPABILITIES ─────────────────────────────────────────────
    add_heading(doc, "5.  AI Agent Capabilities")
    add_body(doc,
        "The AI agent is powered by OpenAI's GPT-4.1-mini model. It is configured "
        "with a system prompt that defines its identity, tone, and the tasks it is "
        "authorised to perform on behalf of WAK Solutions. The table below summarises "
        "what the AI can and cannot do."
    )
    doc.add_paragraph("")
    add_table(doc,
        ["Capability", "Status", "Detail"],
        [
            ["Respond in the customer's language",  "YES — Can Do",  "Detects and replies in Arabic, English, Mandarin Chinese, and other languages automatically."],
            ["Answer product enquiries",            "YES — Can Do",  "Handles questions about products, services, pricing, and operating hours from its knowledge base."],
            ["Look up a specific order",            "YES — Can Do",  "When a customer provides an order number, the AI queries the live database and reports the current status."],
            ["Handle complaints",                   "YES — Can Do",  "Acknowledges the complaint, apologises appropriately, and escalates to a human agent for resolution."],
            ["Escalate to a human agent",           "YES — Can Do",  "Automatically hands off conversations it cannot resolve, with a reason recorded for the agent."],
            ["Maintain conversation history",       "YES — Can Do",  "Remembers the last 20 messages so customers do not have to repeat themselves."],
            ["Process payments or place orders",    "NO — Cannot",   "The AI has read-only access to order data and cannot make changes or process transactions."],
            ["Access personal account details",     "NO — Cannot",   "The AI only looks up orders by order number; it cannot access full account or payment information."],
            ["Operate across other channels",       "NO — Cannot",   "The current implementation is WhatsApp-only; email, web chat, or SMS are not currently connected."],
        ],
        col_widths=[2.4, 1.1, 3.7]
    )

    doc.add_paragraph("")
    add_heading(doc, "Language Support", level=2)
    add_body(doc,
        "The AI does not require manual language selection. It detects the language "
        "of the customer's message and responds accordingly. Supported languages "
        "include Arabic, English, Mandarin Chinese, French, and others supported "
        "natively by GPT-4.1-mini."
    )

    add_page_break(doc)

    # ─── PAGE 6: INFRASTRUCTURE ──────────────────────────────────────────────
    add_heading(doc, "6.  Infrastructure & Integrations")
    add_body(doc,
        "The entire platform runs on cloud infrastructure with no on-premise hardware "
        "required. The two application services are hosted on Railway, a managed "
        "cloud platform, while the database is provided by Neon, a managed PostgreSQL "
        "service. Both are production-grade services with automatic scaling and uptime "
        "guarantees."
    )
    doc.add_paragraph("")
    topo_buf = make_topology_diagram()
    add_image(doc, topo_buf, width_inches=6.2)
    add_caption(doc, "Figure 4 — Deployment Topology")
    doc.add_paragraph("")

    add_heading(doc, "External Services", level=2)
    add_table(doc,
        ["Service", "Provider", "Role", "Managed By"],
        [
            ["WhatsApp Channel",       "Meta (Facebook)",   "Customer messaging channel — sends and receives all WhatsApp messages.", "Meta Cloud"],
            ["AI Language Model",      "OpenAI",            "Understands customer messages and generates contextual replies.",        "OpenAI Cloud"],
            ["Application Hosting",    "Railway",           "Hosts both the AI Service and the Agent Dashboard as separate services.","Railway Cloud"],
            ["Database",               "Neon",              "Stores all messages, escalations, and conversation records.",            "Neon Cloud"],
            ["Push Notifications",     "Web Push / VAPID",  "Delivers browser and mobile notifications to agents.",                  "Self-managed via Railway"],
        ],
        col_widths=[1.8, 1.4, 3.0, 1.8]
    )

    add_page_break(doc)

    # ─── PAGE 7: CONSTRAINTS & NEXT STEPS ────────────────────────────────────
    add_heading(doc, "7.  Assumptions, Constraints & Next Steps")

    add_heading(doc, "Current State (v1.0)", level=2)
    add_body(doc,
        "The current version is a fully functional system. The following items "
        "represent known characteristics of the v1.0 build that should be addressed "
        "before a high-volume production launch:"
    )
    doc.add_paragraph("")
    constraints = [
        ("Authentication",          "Agent login uses a single shared password. For a multi-agent team, individual accounts with role-based access should be introduced."),
        ("Session Storage",         "User sessions are held in server memory, which means agents are logged out if the server restarts. A Redis-backed session store is recommended for production."),
        ("Real-time Updates",       "The dashboard refreshes conversation data every 3–10 seconds by polling. A WebSocket-based push would reduce latency and server load at scale."),
        ("Single WhatsApp Number",  "The platform is configured for one WhatsApp Business number. Multiple numbers or teams would require additional configuration."),
        ("Order Lookup Scope",      "The AI can look up orders by order number but cannot access customer account history or initiate any changes — by design."),
    ]
    add_table(doc,
        ["Area", "Current Behaviour / Constraint"],
        constraints,
        col_widths=[2.0, 5.2]
    )

    doc.add_paragraph("")
    add_heading(doc, "Recommended Next Steps", level=2)
    next_steps = [
        "Introduce individual agent accounts with login audit logging.",
        "Move session storage to a persistent Redis or database-backed store.",
        "Replace polling with WebSockets for real-time message delivery in the dashboard.",
        "Add a reporting module: volume of conversations, escalation rate, average response time.",
        "Extend order lookup to support additional query types (e.g. search by phone number).",
        "Evaluate adding additional channels (e.g. web chat widget) using the same AI backend.",
    ]
    for ns in next_steps:
        add_bullet(doc, ns)

    # ─── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved → {output_path}")


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(__file__), "WAK_Solution_Design.docx")
    build_document(out)
